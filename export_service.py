"""
export_service.py — Xuất tài liệu du lịch chuyên nghiệp v4.0
Fixes:
  ✅ Font DejaVu (sẵn trên hệ thống) hỗ trợ tiếng Việt đầy đủ
  ✅ Parse lịch trình: xử lý đúng ** quanh NGÀY, theme trailing *
  ✅ Parse bảng chi phí: strip markdown trong cells (**, *)
  ✅ DOCX: fix RGBColor, shading, table rendering
  ✅ XLSX: điền đúng data từ bảng markdown
  ✅ Tất cả format xuất đúng nội dung thực tế từ AI response
"""

import io, os, re
from datetime import datetime

# ══════════════════════════════════════════════════════════════
# FONT — DejaVuSans (sẵn trên Ubuntu, hỗ trợ tiếng Việt)
# ══════════════════════════════════════════════════════════════
_FONT_DIR = "/usr/share/fonts/truetype/dejavu"
_FONTS = {
    "reg": os.path.join(_FONT_DIR, "DejaVuSans.ttf"),
    "bold": os.path.join(_FONT_DIR, "DejaVuSans-Bold.ttf"),
    "it": os.path.join(_FONT_DIR, "DejaVuSans-Oblique.ttf"),
    "bi": os.path.join(_FONT_DIR, "DejaVuSans-BoldOblique.ttf"),
}

# Fallback chains
_FALLBACKS = [
    {
        "reg": "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        "bold": "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
        "it": "/usr/share/fonts/truetype/freefont/FreeSansOblique.ttf",
        "bi": "/usr/share/fonts/truetype/freefont/FreeSansBoldOblique.ttf",
    },
    {
        "reg": "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "bold": "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
        "it": "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf",
        "bi": "/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf",
    },
]

_resolved_fonts = None
_fonts_registered = False


def _get_fonts():
    global _resolved_fonts
    if _resolved_fonts:
        return _resolved_fonts
    # Try primary
    if all(os.path.exists(p) for p in _FONTS.values()):
        _resolved_fonts = _FONTS
        return _resolved_fonts
    # Try fallbacks
    for fb in _FALLBACKS:
        if all(os.path.exists(p) for p in fb.values()):
            _resolved_fonts = fb
            return _resolved_fonts
    # Partial fallback: reg+bold only
    for fb in [_FONTS] + _FALLBACKS:
        if os.path.exists(fb["reg"]) and os.path.exists(fb["bold"]):
            _resolved_fonts = {
                "reg": fb["reg"],
                "bold": fb["bold"],
                "it": fb["reg"],
                "bi": fb["bold"],
            }
            return _resolved_fonts
    _resolved_fonts = None
    return None


def _register_fonts():
    global _fonts_registered
    if _fonts_registered:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.pdfmetrics import registerFontFamily

    fonts = _get_fonts()
    if fonts is None:
        _fonts_registered = True
        return
    try:
        if "AriaFont" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("AriaFont", fonts["reg"]))
            pdfmetrics.registerFont(TTFont("AriaFont-B", fonts["bold"]))
            pdfmetrics.registerFont(TTFont("AriaFont-I", fonts["it"]))
            pdfmetrics.registerFont(TTFont("AriaFont-BI", fonts["bi"]))
            registerFontFamily(
                "AriaFont",
                normal="AriaFont",
                bold="AriaFont-B",
                italic="AriaFont-I",
                boldItalic="AriaFont-BI",
            )
    except Exception as e:
        print(f"[FONT] Warning: {e}")
    _fonts_registered = True


def _F(bold=False, italic=False) -> str:
    fonts = _get_fonts()
    if fonts is None:
        if bold and italic:
            return "Helvetica-BoldOblique"
        if bold:
            return "Helvetica-Bold"
        if italic:
            return "Helvetica-Oblique"
        return "Helvetica"
    if bold and italic:
        return "AriaFont-BI"
    if bold:
        return "AriaFont-B"
    if italic:
        return "AriaFont-I"
    return "AriaFont"


# ══════════════════════════════════════════════════════════════
# TEXT / PARSE HELPERS
# ══════════════════════════════════════════════════════════════


def _strip_md(text: str) -> str:
    """Strip markdown formatting, returning clean plain text."""
    if not text:
        return ""
    # Bold+italic, bold, italic
    text = re.sub(r"\*{3}(.+?)\*{3}", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\*{2}(.+?)\*{2}", r"\1", text, flags=re.DOTALL)
    text = re.sub(r"\*([^*\n]+?)\*", r"\1", text)
    # Headers
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    # Code
    text = re.sub(r"`(.+?)`", r"\1", text)
    # Links
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    # Trailing/leading *
    text = text.strip("*").strip()
    return text.strip()


def _extract_duration(text: str) -> str:
    m = re.search(r"(\d+)\s*ngày\s*(\d+)?\s*đêm?", text, re.IGNORECASE)
    if m:
        nights = m.group(2) or str(int(m.group(1)) - 1)
        return f"{m.group(1)} ngày {nights} đêm"
    m = re.search(r"(\d+)[Nn](\d+)[ĐĐđd]", text)
    if m:
        return f"{m.group(1)} ngày {m.group(2)} đêm"
    m = re.search(r"(\d+)\s*(?:days?|ngày)", text, re.IGNORECASE)
    if m:
        return f"{m.group(1)} ngày"
    return ""


def _parse_itinerary_days(text: str) -> list:
    """
    Parse AI response text into list of day dicts.
    Handles: **NGÀY X — Theme**, NGÀY X: Theme, ## Day X, etc.
    """
    # Step 1: Normalize — strip ** wrapping around NGÀY lines
    # Pattern: **NGÀY 1 — ...** → NGÀY 1 — ...
    normalized = re.sub(
        r"\*+\s*((?:NGÀY|Ngày|DAY|Day)\s*\d+[^\n]*?)\s*\*+",
        r"\1",
        text,
        flags=re.IGNORECASE,
    )

    # Step 2: Try primary pattern: NGÀY X followed by rest of line
    day_pat = re.compile(
        r"(?:^|\n)[ \t]*(?:NGÀY|Ngày|DAY|Day)\s*(\d+)\s*[:\-—–]?\s*([^\n]*)",
        re.IGNORECASE,
    )
    matches = list(day_pat.finditer(normalized))

    if matches:
        days = []
        for i, m in enumerate(matches):
            start = m.end()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(normalized)
            theme = _strip_md(m.group(2).strip().strip("—-–:").strip())
            content = normalized[start:end]
            days.append(_build_day(f"Ngày {m.group(1)}", theme, content))
        return days

    # Step 3: Markdown header pattern ## Day X / ### Ngày X
    header_pat = re.compile(
        r"\n#{1,4}\s*(?:Ngày|NGÀY|Day)\s*(\d+)\s*[-—:–]?\s*([^\n]+)", re.IGNORECASE
    )
    matches2 = list(header_pat.finditer(normalized))
    if matches2:
        days = []
        for i, m in enumerate(matches2):
            start = m.end()
            end = matches2[i + 1].start() if i + 1 < len(matches2) else len(normalized)
            theme = _strip_md(m.group(2).strip())
            content = normalized[start:end]
            days.append(_build_day(f"Ngày {m.group(1)}", theme, content))
        return days

    return _fallback_days(normalized)


def _build_day(day_label: str, theme: str, content: str) -> dict:
    """Parse content block into items list."""
    items = []
    lines = content.split("\n")

    # Extract theme from first content line if missing
    if not theme:
        for line in lines:
            s = line.strip()
            if s and not s.startswith(("-", "•", "*", "⏰", "🍽", "🏨", "💡")):
                theme = _strip_md(s)
                break

    EMOJI_MARKERS = [
        "⏰",
        "🍽",
        "🏨",
        "💡",
        "🌅",
        "🌞",
        "🌆",
        "🌙",
        "🚢",
        "🎭",
        "🏖",
        "🏔",
        "🌊",
        "🚣",
        "🎣",
        "🛵",
        "🚗",
        "✈",
        "🛏",
        "🍜",
        "🍺",
        "🎒",
        "🗺",
        "📍",
    ]

    for line in lines:
        line = line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue

        # Skip table rows and separator lines
        if stripped.startswith("|") or re.match(r"^[\|\-\:\s]+$", stripped):
            continue
        # Skip divider lines
        if re.match(r"^[-=]{3,}$", stripped):
            continue

        clean = _strip_md(stripped)
        if not clean:
            continue

        if any(e in stripped for e in EMOJI_MARKERS):
            items.append({"type": "item", "text": clean})
        elif re.match(r"^[-•*]\s+", stripped):
            items.append({"type": "item", "text": _strip_md(stripped[2:].strip())})
        elif re.match(r"^\d+\.\s+", stripped):
            items.append(
                {"type": "item", "text": _strip_md(re.sub(r"^\d+\.\s*", "", stripped))}
            )
        elif re.match(r"^\s{4,}", line):
            items.append({"type": "sub", "text": clean})
        elif stripped:
            items.append({"type": "note", "text": clean})

    return {"day": day_label, "theme": theme, "items": items}


def _fallback_days(text: str) -> list:
    """Last resort: split by blank lines."""
    blocks = re.split(r"\n{3,}", text.strip())
    days = []
    for idx, block in enumerate(blocks[:7], 1):
        if block.strip():
            lines = [_strip_md(l) for l in block.split("\n") if l.strip()]
            days.append(
                {
                    "day": f"Ngày {idx}",
                    "theme": lines[0] if lines else "",
                    "items": [{"type": "item", "text": l} for l in lines[1:] if l],
                }
            )
    return days


def _parse_cost_table(text: str) -> list:
    """
    Parse markdown cost table from AI response.
    Handles ** bold formatting inside cells.
    Returns list of row dicts: {category, budget, mid, luxury}
    """
    rows = []
    seen = set()

    # Find all table lines (start with |, not separator lines)
    table_lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("|") and "|" in stripped[1:]:
            # Skip separator lines like |---|---|
            if not re.match(r"^\|[\s\-:|]+\|", stripped):
                table_lines.append(stripped)

    for line in table_lines:
        # Parse cells
        cells = [_strip_md(c.strip()) for c in line.strip("|").split("|")]
        cells = [c for c in cells if c]
        if not cells:
            continue

        # Skip header row
        first_low = cells[0].lower()
        if any(kw in first_low for kw in ["hạng mục", "category", "tiêu chí", "item"]):
            continue

        key = cells[0]
        if key in seen or not key:
            continue
        seen.add(key)

        row = {"category": key}
        if len(cells) > 1:
            row["budget"] = cells[1]
        if len(cells) > 2:
            row["mid"] = cells[2]
        if len(cells) > 3:
            row["luxury"] = cells[3]
        rows.append(row)

    if rows:
        return rows

    # Fallback: keyword extraction from text
    patterns = [
        (r"(?:vé máy bay|flight)[^\n:]*?:\s*([^\n|,]{3,60})", "Vé máy bay"),
        (
            r"(?:lưu trú|khách sạn|hotel|accommodation)[^\n:]*?:\s*([^\n|,]{3,60})",
            "Lưu trú",
        ),
        (r"(?:ăn uống|food|meals?)[^\n:]*?:\s*([^\n|,]{3,60})", "Ăn uống"),
        (r"(?:tham quan|tour|sightseeing)[^\n:]*?:\s*([^\n|,]{3,60})", "Tham quan"),
        (r"(?:di chuyển|transport)[^\n:]*?:\s*([^\n|,]{3,60})", "Di chuyển"),
        (r"(?:tổng|total)[^\n:]*?:\s*([^\n|,]{3,60})", "TỔNG"),
    ]
    for pattern, label in patterns:
        m = re.search(pattern, text, re.IGNORECASE)
        if m and label not in seen:
            val = _strip_md(m.group(1).strip().split("\n")[0][:80])
            if val:
                rows.append({"category": label, "budget": val})
                seen.add(label)

    return rows


def _extract_tips(text: str) -> list:
    tips = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        low = line.lower()
        if any(
            kw in low
            for kw in [
                "lưu ý",
                "cảnh báo",
                "💡",
                "tip",
                "chú ý",
                "mẹo",
                "kinh nghiệm",
                "cẩn thận",
                "khuyến nghị",
                "ghi chú",
            ]
        ):
            clean = _strip_md(line)
            if 10 < len(clean) < 350:
                tips.append(clean)
    return tips[:12]


# ══════════════════════════════════════════════════════════════
# PDF — Lịch trình du lịch
# ══════════════════════════════════════════════════════════════


def export_itinerary_pdf(
    destination: str, response_text: str, duration: str = "", lang: str = "vi"
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        HRFlowable,
        KeepTogether,
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    _register_fonts()
    if not duration:
        duration = _extract_duration(response_text)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=16 * mm,
        bottomMargin=16 * mm,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
    )
    PAGE_W = A4[0] - 36 * mm

    TEAL = colors.HexColor("#1a3a3a")
    TERRA = colors.HexColor("#c4845a")
    LGREY = colors.HexColor("#f0ede8")
    MGREY = colors.HexColor("#7a6e64")
    WHITE = colors.white
    GOLD = colors.HexColor("#d4b07a")

    def S(name, bold=False, italic=False, **kw):
        return ParagraphStyle(name, fontName=_F(bold, italic), **kw)

    sTitle = S(
        "title",
        bold=True,
        fontSize=22,
        textColor=WHITE,
        alignment=TA_CENTER,
        spaceAfter=3,
    )
    sSub = S("sub", fontSize=10, textColor=GOLD, alignment=TA_CENTER, spaceAfter=2)
    sMeta = S(
        "meta", fontSize=8, textColor=colors.HexColor("#90b0b0"), alignment=TA_CENTER
    )
    sDH = S("dhdr", bold=True, fontSize=11, textColor=WHITE, leftIndent=8)
    sItem = S(
        "item", fontSize=9.5, textColor=TEAL, leftIndent=12, spaceAfter=3, leading=15
    )
    sNote = S(
        "note",
        italic=True,
        fontSize=9,
        textColor=MGREY,
        leftIndent=12,
        spaceAfter=2,
        leading=13,
    )
    sSub2 = S(
        "sub2",
        fontSize=9,
        textColor=colors.HexColor("#2d5555"),
        leftIndent=20,
        spaceAfter=2,
        leading=13,
    )
    sTh = S("th", bold=True, fontSize=9, textColor=WHITE)
    sFoot = S("foot", italic=True, fontSize=7.5, textColor=MGREY, alignment=TA_CENTER)
    sCH = S("ch", bold=True, fontSize=10, textColor=WHITE, leftIndent=8)

    story = []
    now = datetime.now().strftime("%d/%m/%Y")

    # Cover header
    cov_data = [
        [Paragraph("LỊCH TRÌNH DU LỊCH", sSub)],
        [Paragraph(destination.upper(), sTitle)],
    ]
    if duration:
        cov_data.append([Paragraph(duration, sSub)])
    cov_data += [
        [Spacer(1, 5)],
        [Paragraph(f"Tạo bởi Aria Travel AI  ·  {now}", sMeta)],
    ]
    cov = Table(cov_data, colWidths=[PAGE_W])
    cov.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TEAL),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 16),
                ("RIGHTPADDING", (0, 0), (-1, -1), 16),
            ]
        )
    )
    story += [cov, Spacer(1, 7 * mm)]

    # Itinerary days
    days = _parse_itinerary_days(response_text)
    if not days:
        # No day structure found — dump as paragraphs
        for para in response_text.split("\n\n"):
            c = _strip_md(para.strip())
            if c:
                story += [Paragraph(c, sItem), Spacer(1, 3)]
    else:
        for d in days:
            bl = []
            lbl = d["day"]
            if d.get("theme"):
                lbl += f"  —  {d['theme']}"
            hdr = Table([[Paragraph(lbl, sDH)]], colWidths=[PAGE_W])
            hdr.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), TERRA),
                        ("TOPPADDING", (0, 0), (-1, -1), 7),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ]
                )
            )
            bl += [hdr, Spacer(1, 3)]
            for item in d["items"]:
                t = item.get("text", "")
                if not t:
                    continue
                if item["type"] == "note":
                    bl.append(Paragraph(t, sNote))
                elif item["type"] == "sub":
                    bl.append(Paragraph(t, sSub2))
                else:
                    bl.append(Paragraph(t, sItem))
            bl.append(Spacer(1, 5 * mm))
            story.append(KeepTogether(bl))

    # Cost table
    cr = _parse_cost_table(response_text)
    if cr:
        ch = Table([[Paragraph("CHI PHÍ ƯỚC TÍNH", sCH)]], colWidths=[PAGE_W])
        ch.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), TEAL),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        story += [ch, Spacer(1, 4)]

        has_multi = any("mid" in r or "luxury" in r for r in cr)
        if has_multi:
            td = [
                [
                    Paragraph(x, sTh)
                    for x in ["Hạng mục", "Budget", "Tầm trung", "Cao cấp"]
                ]
            ]
            cw = [PAGE_W * 0.36, PAGE_W * 0.21, PAGE_W * 0.21, PAGE_W * 0.22]
        else:
            td = [[Paragraph(x, sTh) for x in ["Hạng mục", "Chi phí"]]]
            cw = [PAGE_W * 0.52, PAGE_W * 0.48]

        for r in cr:
            cat = r.get("category", "")
            is_total = bool(re.search(r"tổng|total", cat, re.IGNORECASE))
            fc = TERRA if is_total else TEAL
            fb = is_total

            def _cell(v, bold=fb, color=fc):
                return Paragraph(
                    v or "—", S("_c", bold=bold, fontSize=9, textColor=color)
                )

            if has_multi:
                td.append(
                    [
                        _cell(cat),
                        _cell(r.get("budget", "—")),
                        _cell(r.get("mid", "—")),
                        _cell(r.get("luxury", "—")),
                    ]
                )
            else:
                td.append([_cell(cat), _cell(r.get("budget", "—"))])

        ct = Table(td, colWidths=cw)
        ct_styles = [
            ("BACKGROUND", (0, 0), (-1, 0), TEAL),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LGREY]),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#ddd")),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("ALIGN", (0, 0), (0, -1), "LEFT"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ]
        for idx, r in enumerate(cr, 1):
            if re.search(r"tổng|total", r.get("category", ""), re.IGNORECASE):
                ct_styles += [
                    ("BACKGROUND", (0, idx), (-1, idx), colors.HexColor("#fff3e8")),
                    ("FONTNAME", (0, idx), (-1, idx), _F(bold=True)),
                ]
        ct.setStyle(TableStyle(ct_styles))
        story += [ct, Spacer(1, 6 * mm)]

    # Footer
    story += [
        HRFlowable(width=PAGE_W, color=colors.HexColor("#e8ddd0"), thickness=0.5),
        Spacer(1, 3),
        Paragraph(
            f"Tạo bởi Aria Travel AI  ·  {now}  ·  Thông tin mang tính chất tham khảo",
            sFoot,
        ),
    ]
    doc.build(story)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════
# PDF — Travel Plan Pro (kế hoạch hoàn chỉnh)
# ══════════════════════════════════════════════════════════════


def export_travel_plan_pdf(
    destination: str, response_text: str, duration: str = "", lang: str = "vi"
) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import mm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        HRFlowable,
        KeepTogether,
    )
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.enums import TA_CENTER

    _register_fonts()
    if not duration:
        duration = _extract_duration(response_text)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
    )
    PAGE_W = A4[0] - 32 * mm

    TEAL = colors.HexColor("#1a3a3a")
    TERRA = colors.HexColor("#c4845a")
    ACC = colors.HexColor("#2d5555")
    GOLD = colors.HexColor("#8B6914")
    LGREY = colors.HexColor("#f0ede8")
    MGREY = colors.HexColor("#7a6e64")
    WHITE = colors.white

    def S(n, bold=False, italic=False, **kw):
        return ParagraphStyle(n, fontName=_F(bold, italic), **kw)

    sT = S(
        "cT", bold=True, fontSize=20, textColor=WHITE, alignment=TA_CENTER, spaceAfter=2
    )
    sS = S("cS", fontSize=10, textColor=colors.HexColor("#d4b07a"), alignment=TA_CENTER)
    sM = S(
        "cM",
        italic=True,
        fontSize=8,
        textColor=colors.HexColor("#90b0b0"),
        alignment=TA_CENTER,
    )
    sSc = S("sc", bold=True, fontSize=11, textColor=WHITE, leftIndent=8)
    sIK = S("ik", bold=True, fontSize=9, textColor=MGREY)
    sIV = S("iv", fontSize=9, textColor=TEAL)
    sDH = S("dh", bold=True, fontSize=10, textColor=WHITE, leftIndent=6)
    sIt = S("it", fontSize=9.5, textColor=TEAL, leftIndent=10, spaceAfter=3, leading=15)
    sNt = S(
        "nt",
        italic=True,
        fontSize=9,
        textColor=MGREY,
        leftIndent=10,
        spaceAfter=2,
        leading=13,
    )
    sSb = S("sb", fontSize=9, textColor=ACC, leftIndent=18, spaceAfter=2, leading=13)
    sTh = S("th", bold=True, fontSize=9, textColor=WHITE)
    sTD = S("td", fontSize=9, textColor=TEAL)
    sTot = S("tot", bold=True, fontSize=9, textColor=TERRA)
    sTip = S(
        "tip",
        italic=True,
        fontSize=9,
        textColor=MGREY,
        leftIndent=8,
        spaceAfter=2,
        leading=13,
    )
    sFt = S("ft", italic=True, fontSize=7, textColor=MGREY, alignment=TA_CENTER)

    story = []
    now = datetime.now().strftime("%d/%m/%Y")

    # Cover
    cov_rows = [
        [Paragraph("KẾ HOẠCH DU LỊCH HOÀN CHỈNH", sS)],
        [Paragraph(destination.upper(), sT)],
        [Paragraph(duration or "Theo kế hoạch", sS)],
        [Spacer(1, 6)],
        [Paragraph(f"Aria Travel AI  ·  Travel Plan Pro  ·  {now}", sM)],
    ]
    cov = Table(cov_rows, colWidths=[PAGE_W])
    cov.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), TEAL),
                ("TOPPADDING", (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 20),
                ("RIGHTPADDING", (0, 0), (-1, -1), 20),
            ]
        )
    )
    story += [cov, Spacer(1, 7 * mm)]

    def _bar(title, color=ACC):
        t = Table([[Paragraph(title, sSc)]], colWidths=[PAGE_W])
        t.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), color),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ]
            )
        )
        return t

    # Section 1: Trip info
    story += [_bar("1. THÔNG TIN CHUYẾN ĐI", TEAL), Spacer(1, 4)]
    info = Table(
        [
            [Paragraph("Điểm đến", sIK), Paragraph(destination, sIV)],
            [Paragraph("Thời gian", sIK), Paragraph(duration or "Theo kế hoạch", sIV)],
            [Paragraph("Tạo ngày", sIK), Paragraph(now, sIV)],
            [Paragraph("Nguồn", sIK), Paragraph("Aria Travel AI", sIV)],
        ],
        colWidths=[PAGE_W * 0.28, PAGE_W * 0.72],
    )
    info.setStyle(
        TableStyle(
            [
                ("ROWBACKGROUNDS", (0, 0), (-1, -1), [WHITE, LGREY]),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#e8ddd0")),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story += [info, Spacer(1, 5 * mm)]

    # Section 2: Detailed itinerary
    story += [_bar("2. LỊCH TRÌNH CHI TIẾT", TERRA), Spacer(1, 4)]
    days = _parse_itinerary_days(response_text)
    if not days:
        for para in response_text.split("\n\n"):
            c = _strip_md(para.strip())
            if c:
                story.append(Paragraph(c, sIt))
    else:
        for d in days:
            bl = []
            lbl = d["day"] + (f"  —  {d['theme']}" if d.get("theme") else "")
            hdr = Table([[Paragraph(lbl, sDH)]], colWidths=[PAGE_W])
            hdr.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), ACC),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ]
                )
            )
            bl += [hdr, Spacer(1, 2)]
            for item in d["items"]:
                t = item.get("text", "")
                if not t:
                    continue
                if item["type"] == "note":
                    bl.append(Paragraph(t, sNt))
                elif item["type"] == "sub":
                    bl.append(Paragraph(t, sSb))
                else:
                    bl.append(Paragraph(t, sIt))
            bl.append(Spacer(1, 4 * mm))
            story.append(KeepTogether(bl))

    # Section 3: Cost table
    cr = _parse_cost_table(response_text)
    if cr:
        story += [_bar("3. BẢNG CHI PHÍ DỰ KIẾN", GOLD), Spacer(1, 4)]
        has_multi = any("mid" in r or "luxury" in r for r in cr)
        if has_multi:
            td = [
                [
                    Paragraph(x, sTh)
                    for x in ["Hạng mục", "Budget", "Tầm trung", "Cao cấp"]
                ]
            ]
            cw = [PAGE_W * 0.36, PAGE_W * 0.21, PAGE_W * 0.21, PAGE_W * 0.22]
        else:
            td = [[Paragraph(x, sTh) for x in ["Hạng mục", "Chi phí ước tính"]]]
            cw = [PAGE_W * 0.5, PAGE_W * 0.5]

        for r in cr:
            cat = r.get("category", "")
            is_total = bool(re.search(r"tổng|total", cat, re.IGNORECASE))
            sty = sTot if is_total else sTD

            def _c(v, s=sty):
                return Paragraph(v or "—", s)

            if has_multi:
                td.append(
                    [
                        _c(cat),
                        _c(r.get("budget", "—")),
                        _c(r.get("mid", "—")),
                        _c(r.get("luxury", "—")),
                    ]
                )
            else:
                td.append([_c(cat), _c(r.get("budget", "—"))])

        ct = Table(td, colWidths=cw)
        cts = [
            ("BACKGROUND", (0, 0), (-1, 0), TEAL),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LGREY]),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#ddd")),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("ALIGN", (0, 0), (0, -1), "LEFT"),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ]
        for idx, r in enumerate(cr, 1):
            if re.search(r"tổng|total", r.get("category", ""), re.IGNORECASE):
                cts += [
                    ("BACKGROUND", (0, idx), (-1, idx), colors.HexColor("#fff3e8")),
                    ("FONTNAME", (0, idx), (-1, idx), _F(bold=True)),
                ]
        ct.setStyle(TableStyle(cts))
        story += [ct, Spacer(1, 5 * mm)]

    # Section 4: Tips
    tips = _extract_tips(response_text)
    if tips:
        story += [
            _bar("4. LƯU Ý & MẸO DU LỊCH", colors.HexColor("#3d6e6e")),
            Spacer(1, 4),
        ]
        for tip in tips:
            story.append(Paragraph(f"• {tip}", sTip))
        story.append(Spacer(1, 5 * mm))

    story += [
        HRFlowable(width=PAGE_W, color=colors.HexColor("#e8ddd0"), thickness=0.5),
        Spacer(1, 3),
        Paragraph(
            f"Tạo bởi Aria Travel AI  ·  {now}  ·  Thông tin mang tính chất tham khảo",
            sFt,
        ),
    ]
    doc.build(story)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════
# DOCX — Lịch trình Word
# ══════════════════════════════════════════════════════════════


def export_itinerary_docx(
    destination: str, response_text: str, duration: str = "", lang: str = "vi"
) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not duration:
        duration = _extract_duration(response_text)

    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    def _rgb(h):
        h = h.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    C_TEAL = _rgb("1a3a3a")
    C_TERRA = _rgb("c4845a")
    C_GOLD = _rgb("b8965a")
    C_GREY = _rgb("7a6e64")
    C_TEAL2 = _rgb("2d5555")
    C_WHITE = _rgb("FFFFFF")

    def _shade_para(para, fill_hex: str):
        """Add background shading to a paragraph."""
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex.lstrip("#").upper())
        pPr.append(shd)

    def _shade_cell(cell, fill_hex: str):
        """Add background shading to a table cell."""
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex.lstrip("#").upper())
        tcPr.append(shd)

    def _add_para(
        text,
        bold=False,
        italic=False,
        size=10,
        color=None,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        indent_cm=0,
        space_before=0,
        space_after=4,
        bg=None,
    ):
        """Add a paragraph with formatting."""
        p = doc.add_paragraph()
        p.alignment = align
        if indent_cm:
            p.paragraph_format.left_indent = Cm(indent_cm)
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        if bg:
            _shade_para(p, bg)
        run = p.add_run(str(text))
        run.bold = bold
        run.italic = italic
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        return p

    now = datetime.now().strftime("%d/%m/%Y")

    # Title
    _add_para(
        f"LỊCH TRÌNH DU LỊCH — {destination.upper()}",
        bold=True,
        size=18,
        color=C_WHITE,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        bg="1a3a3a",
        space_before=8,
        space_after=8,
    )
    if duration:
        _add_para(
            f"{duration}  ·  Tạo bởi Aria Travel AI  ·  {now}",
            italic=True,
            size=10,
            color=C_GOLD,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            space_after=12,
        )
    doc.add_paragraph()

    # Days
    days = _parse_itinerary_days(response_text)
    if not days:
        _add_para(_strip_md(response_text[:5000]), size=10, color=C_TEAL)
    else:
        for d in days:
            lbl = d["day"] + (f"  —  {d['theme']}" if d.get("theme") else "")
            _add_para(
                lbl,
                bold=True,
                size=11,
                color=C_WHITE,
                bg="c4845a",
                space_before=8,
                space_after=4,
            )
            for item in d["items"]:
                t = item.get("text", "")
                if not t:
                    continue
                if item["type"] == "note":
                    _add_para(
                        t,
                        italic=True,
                        size=9.5,
                        color=C_GREY,
                        indent_cm=0.6,
                        space_after=2,
                    )
                elif item["type"] == "sub":
                    _add_para(t, size=9.5, color=C_TEAL2, indent_cm=1.2, space_after=2)
                else:
                    _add_para(t, size=10, color=C_TEAL, indent_cm=0.6, space_after=3)
            doc.add_paragraph()

    # Cost table
    cr = _parse_cost_table(response_text)
    if cr:
        _add_para(
            "BẢNG CHI PHÍ DỰ KIẾN",
            bold=True,
            size=12,
            color=C_WHITE,
            bg="1a3a3a",
            space_before=10,
            space_after=6,
        )
        has_multi = any("mid" in r or "luxury" in r for r in cr)
        n_cols = 4 if has_multi else 2
        headers = (
            ["Hạng mục", "Budget", "Tầm trung", "Cao cấp"]
            if has_multi
            else ["Hạng mục", "Chi phí ước tính"]
        )

        tbl = doc.add_table(rows=1 + len(cr), cols=n_cols)
        tbl.style = "Table Grid"

        # Header row
        for ci, h in enumerate(headers[:n_cols]):
            cell = tbl.rows[0].cells[ci]
            cell.text = ""
            _shade_cell(cell, "1a3a3a")
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = C_WHITE
            run.font.name = "Arial"
            run.font.size = Pt(9)

        # Data rows
        for ridx, row in enumerate(cr, 1):
            is_total = bool(
                re.search(r"tổng|total", row.get("category", ""), re.IGNORECASE)
            )
            bg_hex = "fff3e8" if is_total else ("f0ede8" if ridx % 2 == 0 else "FFFFFF")

            vals = [row.get("category", "")]
            if has_multi:
                vals += [
                    row.get("budget", "—"),
                    row.get("mid", "—"),
                    row.get("luxury", "—"),
                ]
            else:
                vals += [row.get("budget", "—")]

            for ci, v in enumerate(vals[:n_cols]):
                cell = tbl.rows[ridx].cells[ci]
                cell.text = ""
                _shade_cell(cell, bg_hex)
                p = cell.paragraphs[0]
                run = p.add_run(str(v) if v else "—")
                run.bold = is_total
                run.font.name = "Arial"
                run.font.size = Pt(9)
                if is_total:
                    run.font.color.rgb = C_TERRA

        doc.add_paragraph()

    # Tips
    tips = _extract_tips(response_text)
    if tips:
        _add_para(
            "LƯU Ý & MẸO DU LỊCH",
            bold=True,
            size=11,
            color=C_WHITE,
            bg="2d5555",
            space_before=8,
            space_after=6,
        )
        for tip in tips:
            _add_para(
                f"• {tip}",
                italic=True,
                size=9.5,
                color=C_GREY,
                indent_cm=0.4,
                space_after=3,
            )
        doc.add_paragraph()

    # Footer
    _add_para(
        f"Tạo bởi Aria Travel AI  ·  {now}  ·  Thông tin mang tính chất tham khảo",
        italic=True,
        size=8,
        color=C_GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════
# XLSX — Bảng chi phí
# ══════════════════════════════════════════════════════════════


def export_cost_excel(
    destination: str, response_text: str, duration: str = "", lang: str = "vi"
) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    if not duration:
        duration = _extract_duration(response_text)

    wb = Workbook()
    ws = wb.active
    ws.title = f"Chi phí - {destination[:18]}"

    def _fill(hex_color: str) -> PatternFill:
        return PatternFill("solid", fgColor=hex_color.lstrip("#").upper())

    def _font(bold=False, color="1a3a3a", size=10, italic=False) -> Font:
        return Font(
            bold=bold,
            color=color.lstrip("#").upper(),
            size=size,
            italic=italic,
            name="Calibri",
        )

    def _border() -> Border:
        s = Side(style="thin", color="E0D8D0")
        return Border(left=s, right=s, top=s, bottom=s)

    def _align(h="left", v="center", wrap=False) -> Alignment:
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

    now = datetime.now().strftime("%d/%m/%Y")

    # Row 1: Title
    ws.merge_cells("A1:G1")
    ws["A1"] = f"BẢNG CHI PHÍ DU LỊCH — {destination.upper()}"
    ws["A1"].font = _font(bold=True, color="FFFFFF", size=14)
    ws["A1"].fill = _fill("1a3a3a")
    ws["A1"].alignment = _align("center")
    ws.row_dimensions[1].height = 38

    # Row 2: Subtitle
    ws.merge_cells("A2:G2")
    ws["A2"] = f"{duration}  ·  Tạo bởi Aria Travel AI  ·  {now}"
    ws["A2"].font = _font(italic=True, color="b8965a", size=9)
    ws["A2"].fill = _fill("1a3a3a")
    ws["A2"].alignment = _align("center")
    ws.row_dimensions[2].height = 22

    # Row 3: Empty spacer
    ws.append([])

    # Row 4: Headers
    cr = _parse_cost_table(response_text)
    has_multi = any("mid" in r or "luxury" in r for r in cr)

    if has_multi:
        headers = ["Hạng mục", "Mô tả", "Budget", "Tầm trung", "Cao cấp", "Ghi chú", ""]
    else:
        headers = ["Hạng mục", "Mô tả", "Chi phí", "", "", "Ghi chú", ""]

    ws.append(headers)
    for ci in range(1, 7):
        cell = ws.cell(4, ci)
        cell.font = _font(bold=True, color="FFFFFF", size=10)
        cell.fill = _fill("c4845a")
        cell.alignment = _align("center")
        cell.border = _border()
    ws.row_dimensions[4].height = 26

    DATA_START = 5

    if cr:
        # Fill from parsed table data
        for ridx, row in enumerate(cr):
            r_idx = DATA_START + ridx
            cat = row.get("category", "")
            bud = row.get("budget", "")
            mid = row.get("mid", "")
            lux = row.get("luxury", "")
            is_total = bool(re.search(r"tổng|total", cat, re.IGNORECASE))
            bg = "fff3e8" if is_total else ("f0ede8" if ridx % 2 == 0 else "FFFFFF")

            if has_multi:
                vals = [cat, "", bud, mid, lux, ""]
            else:
                vals = [cat, "", bud, "", "", ""]

            for ci, v in enumerate(vals, 1):
                cell = ws.cell(r_idx, ci, v)
                cell.fill = _fill(bg)
                cell.border = _border()
                cell.alignment = _align("center" if ci > 1 else "left", wrap=(ci == 1))
                cell.font = _font(
                    bold=is_total, color="c4845a" if is_total else "1a3a3a", size=9.5
                )
        data_end = DATA_START + len(cr) - 1

        # Add total row if not already present
        last_cat = cr[-1].get("category", "").lower() if cr else ""
        has_total_row = bool(re.search(r"tổng|total", last_cat))
        if not has_total_row:
            tr = data_end + 1
            ws.merge_cells(f"A{tr}:B{tr}")
            ws.cell(tr, 1, "TỔNG CHI PHÍ ƯỚC TÍNH")
            for ci in range(1, 7):
                cell = ws.cell(tr, ci)
                cell.fill = _fill("fff3e8")
                cell.border = _border()
                cell.font = _font(bold=True, color="c4845a", size=10)
                cell.alignment = _align("center" if ci > 1 else "left")
            ws.row_dimensions[tr].height = 26
            total_row = tr
        else:
            total_row = data_end
    else:
        # Default empty template with common categories
        DEFAULTS = [
            ("Vé máy bay khứ hồi", "Cả hai chiều"),
            ("Khách sạn / Homestay", "Chi phí lưu trú / đêm"),
            ("Ăn uống", "3 bữa / ngày"),
            ("Tham quan / Vé vào cửa", "Các điểm tham quan"),
            ("Di chuyển nội địa", "Xe máy, taxi, xe ôm"),
            ("Mua sắm / Quà tặng", "Đặc sản, hàng lưu niệm"),
            ("Dự phòng", "Chi phí phát sinh"),
        ]
        for ridx, (cat, desc) in enumerate(DEFAULTS):
            r_idx = DATA_START + ridx
            ws.cell(r_idx, 1, cat)
            ws.cell(r_idx, 2, desc)
            ws.cell(r_idx, 3, 0)
            ws.cell(r_idx, 4, 0)
            ws.cell(r_idx, 5, 0)
            ws.cell(r_idx, 6, "")
            bg = "f0ede8" if ridx % 2 == 0 else "FFFFFF"
            for ci in range(1, 7):
                cell = ws.cell(r_idx, ci)
                cell.fill = _fill(bg)
                cell.border = _border()
                cell.alignment = _align("center" if ci > 1 else "left")
                cell.font = _font(size=9.5)
                if ci in (3, 4, 5):
                    cell.number_format = "#,##0"
        data_end = DATA_START + len(DEFAULTS) - 1
        tr = data_end + 1
        ws.cell(tr, 1, "TỔNG CHI PHÍ")
        for ci in (3, 4, 5):
            col = get_column_letter(ci)
            ws.cell(tr, ci, f"=SUM({col}{DATA_START}:{col}{data_end})")
            ws.cell(tr, ci).number_format = "#,##0"
        for ci in range(1, 7):
            cell = ws.cell(tr, ci)
            cell.fill = _fill("fff3e8")
            cell.border = _border()
            cell.font = _font(bold=True, color="c4845a", size=11)
            cell.alignment = _align("center" if ci > 1 else "left")
        ws.row_dimensions[tr].height = 26
        total_row = tr

    # Notes section
    ns = total_row + 2
    ws.merge_cells(f"A{ns}:G{ns}")
    ws.cell(ns, 1, "GHI CHÚ & LƯU Ý")
    ws.cell(ns, 1).font = _font(bold=True, color="FFFFFF", size=10)
    ws.cell(ns, 1).fill = _fill("2d5555")
    ws.cell(ns, 1).alignment = _align("left")
    ws.row_dimensions[ns].height = 24

    notes = [
        "• Chi phí trên là ước tính tham khảo, có thể thay đổi theo mùa và nhà cung cấp.",
        "• Nên đặt phòng và vé sớm ít nhất 2–4 tuần để có giá tốt nhất.",
        "• Budget: tìm Grab / xe ôm / hostel thay vì taxi / khách sạn.",
        "• Tổng chi phí chưa bao gồm bảo hiểm du lịch (~100–200k/chuyến).",
        f"• Điểm đến: {destination}  |  Thời gian: {duration or 'Chưa xác định'}",
    ]
    # Also add any extracted tips from response
    tips = _extract_tips(response_text)
    for tip in tips[:3]:
        notes.append(f"• {tip}")

    for nidx, note in enumerate(notes):
        nr = ns + 1 + nidx
        ws.merge_cells(f"A{nr}:G{nr}")
        ws.cell(nr, 1, note)
        ws.cell(nr, 1).font = _font(size=9, italic=True, color="555555")
        ws.cell(nr, 1).alignment = _align("left", wrap=True)
        ws.row_dimensions[nr].height = 18

    # Column widths
    for ci, w in enumerate([30, 28, 18, 18, 18, 22, 0], 1):
        if w:
            ws.column_dimensions[get_column_letter(ci)].width = w

    ws.freeze_panes = "A5"

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()
